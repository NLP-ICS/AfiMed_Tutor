"""Generate report/final_report.docx using python-docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# ── Styles ────────────────────────────────────────────────────────────────────
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)

for hname, size, bold in [
    ("Heading 1", 14, True),
    ("Heading 2", 12, True),
    ("Heading 3", 11, True),
]:
    style = doc.styles[hname]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text, bold=False, italic=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for para in cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AfriMed Tutor: A Guideline-Grounded Study Assistant\nfor African Medical Students")
r.bold = True
r.font.size = Pt(16)
r.font.name = "Times New Roman"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("ICS4554 Natural Language Processing — Final Project Report\nAshesi University · Spring 2026\n[insert GitHub repository URL]")
r2.font.size = Pt(11)
r2.font.name = "Times New Roman"
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM IDENTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Problem Identification", 1)

add_para(
    "Sub-Saharan Africa faces a severe and widening gap in clinical expertise. The region carries "
    "25% of the global disease burden while accounting for fewer than 3% of the world's health workers, "
    "and the majority of practicing clinicians received their training in systems where access to "
    "high-quality, up-to-date clinical decision support is limited (WHO, 2023). Medical students studying "
    "in this environment often rely on textbooks that pre-date current African clinical guidelines, or on "
    "international resources calibrated to disease prevalence patterns and drug formularies that differ "
    "materially from those encountered in African practice."
)

add_para(
    "A recent benchmark study by Olatunji et al. (2025) quantified this gap directly. The AfriMed-QA "
    "dataset contains 3,724 multiple-choice questions drawn from actual African medical licensing and "
    "postgraduate examinations. When state-of-the-art frontier LLMs are evaluated on this benchmark, "
    "accuracy ranges from 57% to 76%, with the largest deficits in Infectious Disease and Obstetrics and "
    "Gynecology: precisely the specialties most distinctive to the African disease burden. These models "
    "were not trained to reason from African clinical guidelines; they were trained predominantly on "
    "English-language text weighted toward North American and European clinical practice."
)

add_para(
    "This project builds AfriMed Tutor, a Retrieval-Augmented Generation (RAG) study assistant that "
    "grounds its responses in a curated corpus of five current African clinical guideline documents "
    "(South Africa PHC and Hospital STGs, Kenya Clinical Guidelines, Ghana STG, WHO IMCI). The system's "
    "design is motivated by two research questions: (1) Does providing African guideline context at "
    "inference time improve a frontier LLM's performance on African clinical MCQs? (2) How does retrieval "
    "method (dense semantic versus sparse keyword) affect system accuracy and explanation quality?"
)

add_para(
    "The system is designed explicitly as a study aid, not a clinical decision tool. Its value is "
    "pedagogical: it helps students understand the clinical reasoning embedded in the guidelines their "
    "future patients will be managed by."
)


# ══════════════════════════════════════════════════════════════════════════════
# 2. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Literature Review", 1)

add_heading("2.1 Medical NLP benchmarks and the African context", 2)

add_para(
    "The AfriMed-QA benchmark (Olatunji et al., 2025) is the primary evaluation dataset for this work. "
    "It contains 3,724 MCQs and 1,236 short-answer questions from medical licensing examinations in Ghana, "
    "Nigeria, Kenya, and South Africa. The paper's key finding, that even the strongest frontier models "
    "underperform on African medical questions, motivates the retrieval-augmented approach taken here. "
    "Notably, AfriMed-QA's MCQ split skews toward Nigerian licensing examinations, a limitation "
    "acknowledged by the authors and relevant to interpreting per-specialty results."
)

add_para(
    "Singhal et al. (2023) introduced Med-PaLM 2, a fine-tuned LLM that achieved expert-level "
    "performance (86.5%) on the US Medical Licensing Examination. While this demonstrates the ceiling "
    "achievable with fine-tuning on curated medical data, it also illustrates the access barrier: "
    "Med-PaLM's training required proprietary curated datasets and large-scale compute inaccessible to "
    "most African research institutions. Our work deliberately avoids fine-tuning and relies entirely on "
    "retrieval at inference time, keeping the system deployable with a single API key."
)

add_para(
    "Fallahpour et al. (2025) (MedRax) extended RAG-based approaches to radiology report generation, "
    "reporting that retrieval augmentation can substitute effectively for domain-specific fine-tuning "
    "when the retrieval corpus is sufficiently targeted. This supports the core hypothesis of our project."
)

add_heading("2.2 Retrieval-Augmented Generation", 2)

add_para(
    "Lewis et al. (2020) introduced the RAG framework, showing that conditioning a generative model on "
    "dynamically retrieved documents at inference time outperforms parametric-only approaches on "
    "knowledge-intensive NLP tasks. The two retrieval paradigms used in this project, dense (FAISS over "
    "normalized embeddings) and sparse (BM25), correspond directly to the retriever families analyzed in "
    "subsequent RAG literature."
)

add_para(
    "Karpukhin et al. (2020) established that dense passage retrieval outperforms BM25 on open-domain QA "
    "tasks that require semantic matching, while Robertson and Zaragoza (2009) showed BM25's continued "
    "strength on terminology-heavy domains. Medical text is terminology-heavy, which partly explains the "
    "relative performance we observe (see §3.4)."
)

add_heading("2.3 Educational NLP in African contexts", 2)

add_para(
    "Boateng et al. (2024) built Kwame for Science, an AI question-answering system for Ghanaian "
    "secondary school students grounded in the national science curriculum. The design parallels ours: "
    "curriculum-grounded retrieval, frontier LLM generation, and explicit citation of source material. "
    "Their finding that grounding in local educational materials improved both accuracy and student trust "
    "directly informed our decision to use African-sourced guidelines rather than international medical corpora."
)

add_heading("2.4 Automated evaluation of NLP systems", 2)

add_para(
    "The LLM-as-judge evaluation protocol used in §3.5 follows the methodology introduced by Zheng et al. "
    "(2023) for MT-Bench and adopted widely in subsequent work. The key methodological concern, that a "
    "model judging its own outputs suffers from same-model bias, led us to designate a different LLM "
    "provider as judge. In our case, the absence of an OpenAI API key required the fallback of using the "
    "same model family as judge, a limitation discussed further in §3.5."
)

add_heading("2.5 Ethics and representation in medical AI", 2)

add_para(
    "Bender et al. (2021) warned of stochastic parrots: language models that reproduce statistical "
    "patterns in training data without genuine comprehension, with harms amplified in high-stakes domains. "
    "In medical AI, this concern is compounded by the geographic and demographic skew of training corpora "
    "identified by Obermeyer et al. (2019). Our explicit grounding in African guidelines is a partial "
    "mitigation, but as §4 discusses, it does not eliminate the risk."
)

add_para(
    "Hu et al. (2025), in a JMIR scoping review of LLM-based health education tools, found consistent "
    "reporting gaps around transparency, bias auditing, and safety disclaimers. AfriMed Tutor addresses "
    "this through mandatory UI disclaimers and the citation-obligated prompt design: every factual claim "
    "is traceable to a retrieved excerpt from a named guideline."
)


# ══════════════════════════════════════════════════════════════════════════════
# 3. SOLUTION AND EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Solution and Evaluation", 1)

add_heading("3.1 System architecture", 2)

add_para(
    "AfriMed Tutor is a Streamlit web application with three interaction modes. A single orchestrator "
    "module (tutor/orchestrator.py) handles mode routing and prompt composition; one LLM call is made "
    "per user interaction. The retrieval pipeline consists of two independent retrievers: a dense "
    "FAISS-based retriever and a BM25 sparse retriever, both indexed over the same 5,633 guideline chunks."
)

add_heading("3.2 Corpus and chunking", 2)

add_para(
    "The guideline corpus comprises five documents totaling approximately 55 MB of PDFs. Text extraction "
    "uses pypdf as the primary engine, with pdfplumber as a fallback. The structure-aware chunker applies "
    "document-family-specific regex patterns to detect section headers before splitting. For example, "
    "Ghana STG chapters follow the pattern of condition names in all-caps with Diagnosis, Investigations, "
    "Management, and Referral subsections; the chunker splits at these subsection boundaries before "
    "applying a secondary word-count split for chunks exceeding 800 tokens, with a 100-token overlap. "
    "Each chunk carries metadata: source_doc, section_title, condition, page_number, and chunk_id. "
    "Mean chunk size is 145 words."
)

add_table(
    headers=["Document", "Source", "Chunks"],
    rows=[
        ["SA Standard Treatment Guidelines, PHC, 8th ed (2024)", "National DoH, South Africa", "2,004"],
        ["SA Standard Treatment Guidelines, Hospital Adults, 6th ed (2024)", "SAHIVSOC, South Africa", "2,103"],
        ["Kenya Clinical Guidelines, 3rd ed", "WHO/Kenya MoH", "794"],
        ["Ghana Standard Treatment Guidelines, 7th ed (2017)", "Ghana MoH", "572"],
        ["WHO IMCI Chart Booklet (2014)", "WHO", "160"],
        ["Total", "", "5,633"],
    ],
    col_widths=[2.8, 1.8, 0.9],
)

add_heading("3.3 Embedding and retrieval", 2)

add_para(
    "Dense retrieval uses the all-MiniLM-L6-v2 sentence-transformer (384-dimensional embeddings) indexed "
    "in a FAISS IndexFlatIP over L2-normalized vectors, implementing cosine similarity via inner product. "
    "This model was chosen over OpenAI text-embedding-3-small or Voyage voyage-3-lite for cost and "
    "offline-reproducibility reasons: all-MiniLM-L6-v2 runs locally without an API key and shows "
    "competitive performance on biomedical retrieval tasks in the MTEB benchmark. The full corpus of "
    "5,633 chunks was embedded in a single batch; the 3,624 MCQ pool questions were also pre-embedded "
    "for the related-questions retriever."
)

add_para(
    "Sparse retrieval uses BM25 (via rank_bm25) over the same chunk texts, with lowercasing and "
    "whitespace tokenization. Both retrievers return the top-k=5 chunks by default. A cosine similarity "
    "threshold of 0.30 was applied after exploratory inspection on 20 held-out queries: below this "
    "threshold, chunks were qualitatively unrelated to the query."
)

add_heading("3.4 Quantitative evaluation: MCQ accuracy", 2)

add_para(
    "We evaluated Claude Sonnet 4.6 on 100 held-out AfriMed-QA MCQs, stratified across five specialties "
    "(20 per specialty), under three conditions: (1) baseline with no retrieval, (2) RAG with dense "
    "retrieval, and (3) RAG with sparse BM25 retrieval. All conditions used the same model and temperature "
    "0.0. Bootstrap 95% confidence intervals were computed with 1,000 resamples."
)

add_para("Table 1: Overall accuracy by condition", bold=True)
add_table(
    headers=["Condition", "Accuracy", "95% CI"],
    rows=[
        ["Baseline (LLM only, no retrieval)", "0.810", "[0.730, 0.880]"],
        ["RAG with dense retrieval", "0.700", "[0.600, 0.790]"],
        ["RAG with sparse BM25 retrieval", "0.760", "[0.670, 0.840]"],
    ],
    col_widths=[3.0, 1.0, 1.5],
)

add_para("Table 2: Per-specialty accuracy", bold=True)
add_table(
    headers=["Specialty", "Baseline", "RAG Dense", "RAG Sparse"],
    rows=[
        ["Surgery", "0.900", "0.800", "0.900"],
        ["Pediatrics", "0.850", "0.850", "0.800"],
        ["Internal Medicine", "0.800", "0.700", "0.850"],
        ["Obstetrics and Gynecology", "0.800", "0.650", "0.700"],
        ["Infectious Disease", "0.700", "0.500", "0.550"],
    ],
    col_widths=[2.2, 1.1, 1.1, 1.1],
)

add_para(
    "Interpreting the RAG underperformance. The baseline (no retrieval) outperforms both RAG conditions "
    "overall. Several factors likely contribute. First, guideline specificity conflicts with AfriMed-QA "
    "answer keys: the MCQs originate largely from Nigerian licensing examinations that may follow "
    "international (WHO, British) clinical guidelines rather than the country-specific protocols in our "
    "corpus. Second, retrieval noise forces the model to guess from incomplete context when the top-5 "
    "chunks do not contain the relevant protocol. Third, BM25 consistently outperforms dense retrieval "
    "(+6 percentage points overall), especially in Internal Medicine (+15 pp) and Surgery (+10 pp), "
    "because medical MCQs contain precise terminology that BM25 matches exactly, while dense retrieval "
    "sometimes retrieves topically related but lexically mismatched excerpts."
)

add_heading("3.5 Qualitative evaluation: LLM-as-judge groundedness", 2)

add_para(
    "We randomly sampled 30 items from the MCQ pool, stratified by specialty, and ran Quiz mode on each "
    "by submitting the gold answer as the student's answer to obtain non-trivial explanations. A judge "
    "LLM scored each explanation on three dimensions using a 0 to 2 scale."
)

add_para(
    "Note on judge configuration. The spec requires the judge to be a different model from the answerer "
    "to reduce same-model bias. As an OpenAI API key was not available, the judge was run with the same "
    "model family (Claude Sonnet 4.6). Results should be interpreted with this caveat.",
    italic=True
)

add_para("Table 3: Mean judge scores (n=29 successfully evaluated)", bold=True)
add_table(
    headers=["Dimension", "Mean Score (out of 2)"],
    rows=[
        ["Groundedness (claims supported by retrieved excerpts)", "0.17"],
        ["Citation accuracy (citations point to real supporting chunks)", "1.45"],
        ["Consistency with gold rationale", "1.83"],
    ],
    col_widths=[4.0, 1.5],
)

add_para(
    "The low groundedness score (0.17/2) is the most striking finding. The judge consistently flagged "
    "that while the model's explanations were medically accurate and consistent with the gold rationale "
    "(1.83/2), the specific factual claims went substantially beyond what the retrieved guideline excerpts "
    "contained. This is expected behavior for a model instructed to produce a 4 to 8 sentence educational "
    "explanation when it has both the gold rationale and retrieved chunks available: the model's "
    "parametric knowledge fills gaps where retrieval falls short. The citation accuracy score (1.45/2) "
    "indicates that when citations are produced, they are mostly correct, but the model does not always "
    "produce citations for claims drawn from parametric memory."
)

add_heading("3.6 Related-question retrieval", 2)

add_para(
    "After each interaction, the system surfaces 3 related practice questions from the AfriMed-QA pool "
    "using cosine similarity over pre-computed question embeddings (Method A). A BM25-based keyword "
    "retriever filtered by specialty is available as Method B. On 10 qualitative inspection cases, "
    "Method A (semantic) produced more topically diverse related questions covering different clinical "
    "angles; Method B tended to retrieve near-duplicate questions with overlapping keywords. For a study "
    "aid, Method A's diversity is preferred."
)

add_heading("3.7 Latency and cost", 2)

add_para("Table 4: Latency and token counts per condition (MCQ evaluation, n=100)", bold=True)
add_table(
    headers=["Condition", "Mean input tokens", "Mean latency (ms)", "p95 latency (ms)"],
    rows=[
        ["Baseline", "152", "2,366", "3,361"],
        ["RAG Dense", "1,675", "4,449", "~7,100"],
        ["RAG Sparse", "1,675", "3,727", "~6,500"],
    ],
    col_widths=[1.8, 1.4, 1.4, 1.4],
)

add_para(
    "Note: MCQ evaluation uses max_tokens=10 to extract a single letter answer. Real Quiz interactions "
    "generate 200 to 400 token explanations. Estimated cost per 100 real Quiz interactions using "
    "Claude Sonnet 4.6 at $3.00 per million input tokens and $15.00 per million output tokens: "
    "RAG input ~$0.50, explanation output ~$0.45, local embedding $0.00. Total: approximately $0.95 "
    "per 100 interactions, which is feasible for institutional deployment."
)


# ══════════════════════════════════════════════════════════════════════════════
# 4. ETHICAL ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Ethical Analysis", 1)

add_heading("4.1 Hallucination risk in clinical contexts", 2)

add_para(
    "The groundedness evaluation (§3.5) confirmed that AfriMed Tutor's explanations routinely include "
    "claims not strictly supported by the retrieved excerpts. In a medical education context this carries "
    "specific risks: a student who treats an uncited claim as authoritative guideline content may mislearn "
    "a drug dose, a contraindication, or a diagnostic criterion. The citation prompt design "
    "([Source: doc, Section: section]) partially mitigates this by making it visible when a claim is "
    "unsupported, but the judge evaluation indicates the model frequently omits citations even when it "
    "extends beyond retrieved context. The mandatory disclaimer shown at all times in the UI "
    "(Study aid only, not for clinical decision-making) is the primary guardrail."
)

add_heading("4.2 Over-reliance and deskilling", 2)

add_para(
    "Pedagogical AI tools risk producing students who can retrieve correct answers from AI systems but "
    "cannot reason independently. AfriMed Tutor's Quiz mode forces active engagement: the student must "
    "select an answer before seeing the explanation. The related-questions feature encourages continued "
    "practice across topics rather than repeated querying of the same item. These are weak safeguards "
    "and the system should not be used as a substitute for clinical reasoning practice under supervision."
)

add_heading("4.3 Geographic and demographic representation bias", 2)

add_para(
    "The AfriMed-QA benchmark skews toward Nigerian clinical contexts. The guideline corpus "
    "overrepresents South African clinical protocols (4,107 of 5,633 chunks, or 73%), which apply to "
    "a health system with relatively high resources compared to much of sub-Saharan Africa. A medical "
    "student in rural Ghana or Uganda will encounter different drug availability, diagnostic capacity, "
    "and epidemiological profiles than what our corpus primarily reflects. Pretraining bias compounds "
    "this: Claude Sonnet 4.6 was trained on data heavily weighted toward North American and European "
    "biomedical literature. Even with African guideline retrieval, the model's implicit prior for "
    "clinical reasoning is non-African. This is a structural limitation of the API-only approach "
    "mandated by the project scope."
)

add_heading("4.4 Data licensing and consent", 2)

add_para(
    "All guideline documents used are either government publications freely distributable for "
    "educational use, or carry explicit non-commercial/educational-use permissions (SA STG, WHO IMCI "
    "CC BY-NC-SA). The AfriMed-QA dataset is CC BY-NC-SA 4.0. The system is not deployed publicly, "
    "has no user data retention, and collects no personal information from students. No IRB process "
    "was required."
)

add_heading("4.5 Equity of access", 2)

add_para(
    "The system currently requires a paid Anthropic API key to operate. This creates an access barrier "
    "that contradicts the equity-of-access motivation for the project. Future work should explore smaller "
    "open-weight models (Llama 3.1 8B, Mistral 7B) that could run locally on consumer hardware, removing "
    "the API cost barrier for institutions with unreliable internet connectivity."
)


# ══════════════════════════════════════════════════════════════════════════════
# 5. TEAM CONTRIBUTIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Team Contributions", 1)

add_table(
    headers=["Component", "Primary contributor", "Supporting contributor"],
    rows=[
        ["Corpus assembly and chunking (corpus/build_corpus.py)", "", ""],
        ["Dense and sparse retrieval (tutor/retriever.py)", "", ""],
        ["AfriMed-QA data loader (data/load_afrimedqa.py)", "", ""],
        ["Orchestrator and prompts (tutor/orchestrator.py, tutor/prompts.py)", "", ""],
        ["Streamlit UI (app.py)", "", ""],
        ["MCQ and retriever evaluation (eval/run_mcq_eval.py, eval/run_retriever_comparison.py)", "", ""],
        ["LLM-as-judge evaluation (eval/run_groundedness_judge.py)", "", ""],
        ["Report writing", "", ""],
    ],
    col_widths=[3.2, 1.4, 1.4],
)


# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading("References", 1)

refs = [
    "Bender, E. M., Gebru, T., McMillan-Major, A., and Shmitchell, S. (2021). On the dangers of stochastic parrots: Can language models be too big? FAccT 2021.",
    "Boateng, G., et al. (2024). Kwame for Science: An AI teaching assistant for science education in Ghana. NeurIPS 2024 Workshop on AI for Education.",
    "Fallahpour, A., et al. (2025). MedRax: Retrieval-augmented generation for radiology report synthesis. arXiv preprint.",
    "Hu, Y., et al. (2025). Large language models for health education: A scoping review. JMIR Medical Education.",
    "Karpukhin, V., et al. (2020). Dense passage retrieval for open-domain question answering. EMNLP 2020.",
    "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. NeurIPS 2020.",
    "Obermeyer, Z., et al. (2019). Dissecting racial bias in an algorithm used to manage the health of populations. Science, 366(6464), 447-453.",
    "Olatunji, T., et al. (2025). AfriMedQA: A benchmark for evaluating large language models on African medical knowledge. arXiv preprint.",
    "Robertson, S., and Zaragoza, H. (2009). The probabilistic relevance framework: BM25 and beyond. Foundations and Trends in Information Retrieval.",
    "Singhal, K., et al. (2023). Large language models encode clinical knowledge. Nature, 620, 172-180.",
    "World Health Organization. (2023). Health workforce: Key facts. WHO.",
    "Zheng, L., et al. (2023). Judging LLM-as-a-judge with MT-Bench and Chatbot Arena. NeurIPS 2023.",
]

for ref in refs:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"


# ── Save ──────────────────────────────────────────────────────────────────────
out = "report/final_report.docx"
doc.save(out)
print("Saved:", out)
